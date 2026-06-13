"""UROP 보고서 .docx 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table(doc, headers, rows, col_widths=None):
    """서식화된 표 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(cell, 'D9E2F3')

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def main():
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)

    # ===== 제목 =====
    title = doc.add_heading('2026년도 1학기 UROP 보고서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 요약문 =====
    doc.add_heading('요약문', level=1)

    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('지원 연구실명', ''),
        ('지원자 학번', ''),
        ('지도교수명', '김상철'),
        ('지원자 성명', ''),
        ('멘토대학원생 성명', ''),
        ('연락처(멘토)', ''),
        ('지원자 연락처', ''),
    ]
    for i, (key, val) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = val
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if info_table.rows[i].cells[0].paragraphs[0].runs else False
        info_table.rows[i].cells[0].width = Cm(5)
        info_table.rows[i].cells[1].width = Cm(11)

    doc.add_paragraph()

    # ===== 지원 동기 =====
    doc.add_heading('지원 동기', level=1)
    doc.add_paragraph(
        '평소 컴퓨터 비전 분야에 관심이 있었고, 해당 연구실에서 UAV를 직접 탐지하는 연구를 '
        '진행하고 있다고 알게되었습니다. 해당 분야에 좀 더 깊은 지식을 쌓고자 연구를 희망하게 '
        '되었고, 이에 지원하게 되었습니다.'
    )

    # ===== 연구 주제 =====
    doc.add_heading('연구 주제', level=1)
    doc.add_paragraph(
        'Anti-UAV 환경에서 UNet 세그멘테이션과 칼만 필터를 융합한 드론 위치 추적 시스템 구현'
    ).bold = True
    doc.add_paragraph(
        '단순히 드론을 박스 형태로 탐지하고 등속도 운동으로 가정하는 기존 방식의 한계를 벗어나, '
        '픽셀 단위로 드론의 형태를 분리해 내고 다양한 모션 모델(등속도, 등가속도, 비선형 CTRV)을 '
        '적용하여 추적 안정성과 정확도를 비교 분석하였다.'
    )

    # ===== 연구 결과 요약 =====
    doc.add_heading('연구 결과 요약', level=1)

    # --- 1. 연구 목표 ---
    doc.add_heading('1. 연구 목표', level=2)
    doc.add_paragraph(
        'UNet 세그멘테이션으로 프레임별 드론 마스크를 생성하고, 중심점(centroid)을 추출하여 '
        '칼만 필터로 시간축 평활화를 적용함으로써 추적 안정성을 개선한다. '
        '세 가지 모션 모델(선형 CV, 등가속도 CA, 비선형 EKF-CTRV)의 성능을 비교 평가한다.'
    )

    # --- 2. 시스템 구조 ---
    doc.add_heading('2. 시스템 구조', level=2)
    doc.add_paragraph(
        '입력 프레임(480×480) → [UNet 세그멘테이션] → 객체 마스크 → '
        '[중심점 추출] → 측정값 z=[x,y] → [칼만 필터 Predict/Update] → 평활화된 위치 추정'
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('UNet: ').bold = True
    p.add_run('VanillaUNet (Encoder-Decoder, depth=4, BatchNorm+ReLU, BCEWithLogitsLoss, AdamW lr=1e-4)')
    p = doc.add_paragraph()
    p.add_run('칼만 필터: ').bold = True
    p.add_run('예측(Predict) → 업데이트(Update) 반복으로 프레임 간 노이즈 제거')

    # --- 3. 모션 모델 ---
    doc.add_heading('3. 비교한 모션 모델', level=2)

    doc.add_paragraph('(1) Constant Velocity (CV) — 선형 칼만 필터', style='List Number')
    doc.add_paragraph('    상태: [x, y, vx, vy] (4차원). 속도가 일정하다고 가정.')
    doc.add_paragraph('(2) Constant Acceleration (CA)', style='List Number')
    doc.add_paragraph('    상태: [x, y, vx, vy, ax, ay] (6차원). 가속도까지 추정.')
    doc.add_paragraph('(3) CTRV — Extended Kalman Filter (EKF)', style='List Number')
    doc.add_paragraph('    상태: [x, y, v, θ, ω] (5차원). 비선형 상태 전이, 야코비안 기반 공분산 전파.')

    # --- 4. 실험 환경 ---
    doc.add_heading('4. 실험 환경', level=2)

    add_table(doc,
        ['항목', '내용'],
        [
            ['데이터셋', 'ANTI-UAV RGBT (적외선/가시광)'],
            ['학습 시퀀스', '4'],
            ['테스트 시퀀스', '1 (1,708 프레임)'],
            ['이미지 크기', '480 × 480'],
            ['학습 Epoch', '50'],
            ['손실 함수', 'BCEWithLogitsLoss'],
            ['Optimizer', 'AdamW (lr=1e-4)'],
        ],
        col_widths=[5, 11]
    )

    doc.add_paragraph()

    # --- 5. 실험 결과 ---
    doc.add_heading('5. 실험 결과', level=2)

    # 5.1 세그멘테이션
    doc.add_heading('5.1 세그멘테이션 성능', level=3)

    add_table(doc,
        ['모델', 'mIoU', 'Dice'],
        [
            ['VanillaUNet', '0.866', '0.926'],
        ],
        col_widths=[6, 5, 5]
    )
    doc.add_paragraph()

    # 5.2 필터 비교
    doc.add_heading('5.2 추적 필터 비교 (Q=0.3, R=0.5)', level=3)

    add_table(doc,
        ['방법', 'CLE (px) ↓', 'Jitter ↓', 'Jitter 감소율', 'Smoothness ↑'],
        [
            ['Baseline (UNet only)', '2.88', '3.63', '—', '1.00'],
            ['Linear KF (등속도, CV)', '3.22', '2.75', '-24.0%', '1.32'],
            ['CA (등가속도)', '3.13', '3.08', '-15.0%', '1.18'],
            ['EKF (CTRV)', '3.27', '3.35', '-7.5%', '1.08'],
        ],
        col_widths=[5.5, 2.8, 2.5, 3, 3]
    )
    doc.add_paragraph()

    # 지표 설명
    p = doc.add_paragraph()
    p.add_run('CLE').bold = True
    p.add_run(': 예측 중심점과 GT 중심점의 유클리드 거리 (낮을수록 정확)')
    p = doc.add_paragraph()
    p.add_run('Jitter').bold = True
    p.add_run(': 프레임 간 추적 변동성 (낮을수록 안정적)')
    p = doc.add_paragraph()
    p.add_run('Smoothness Ratio').bold = True
    p.add_run(': 필터 적용 전/후 안정성 비율 (>1이면 개선)')

    doc.add_paragraph()

    # 5.3 Q 파라미터
    doc.add_heading('5.3 Q 파라미터 민감도 분석 (Linear KF)', level=3)

    add_table(doc,
        ['Q (process noise)', 'CLE 변화', 'Jitter 감소율', '특성'],
        [
            ['0.01', '+2.67 px', '54%', '최대 평활화, 위치 lag 발생'],
            ['0.1', '+0.79 px', '34%', '균형 후보'],
            ['0.3 (채택)', '+0.34 px', '24%', 'CLE 손해 최소 + 유의미한 평활화'],
            ['1.0', '+0.10 px', '14%', '위치 정확, 평활화 미미'],
        ],
        col_widths=[3.5, 2.5, 3, 7]
    )
    doc.add_paragraph()

    # --- 6. 결론 ---
    doc.add_heading('6. 결론', level=2)

    conclusions = [
        'UNet 세그멘테이션은 mIoU 0.866, Dice 0.926의 높은 정확도를 달성하였다.',
        '선형 칼만 필터(등속도 모델)가 종합적으로 가장 효과적이며, 추적 안정성을 24% 개선하면서 위치 정확도 손실은 0.34px에 불과하였다.',
        '등가속도(CA) 모델은 위치 정확도에서 소폭 우위를 보였으나, 안정성 개선은 제한적이었다.',
        'EKF(CTRV)는 비선형 모델의 상태 차원(5D)과 관측 차원(2D)의 불균형으로 인해 본 데이터셋의 등속 위주 운동에서는 효과가 미미하였다.',
        '칼만 필터의 핵심 기여는 위치 정확도 향상보다 추적 궤적의 시간적 안정성 확보에 있으며, Q 파라미터를 통해 정확도-안정성 간 trade-off를 제어할 수 있다.',
    ]
    for c in conclusions:
        doc.add_paragraph(c, style='List Number')

    doc.add_paragraph()

    # --- 7. 향후 연구 ---
    doc.add_heading('7. 향후 연구 방향', level=2)

    future = [
        '급선회/고기동 시퀀스에서 EKF 재평가',
        '가림(occlusion) 상황에서 칼만 필터의 예측 기반 추적 유지 검증',
        'IMM(Interacting Multiple Model) 필터로 모션 모드 자동 전환',
        '다중 객체 추적 확장',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== 서명 =====
    doc.add_paragraph('─' * 40)
    doc.add_paragraph()
    p = doc.add_paragraph(
        'UROP 수업으로 위 학생이 지원한 해당 연구실에서 2026년도 1학기 동안 '
        '근무하고자 신청서를 제출합니다.'
    )
    doc.add_paragraph()
    doc.add_paragraph('2026년          월                    일')
    doc.add_paragraph()
    doc.add_paragraph('신청자:                                          (서명)')
    doc.add_paragraph()
    doc.add_paragraph('UROP 지도교수:                                   (서명)')

    # 저장
    output_path = 'UROP_보고서.docx'
    doc.save(output_path)
    print(f'✓ 보고서 생성 완료: {output_path}')


if __name__ == '__main__':
    main()
